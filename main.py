"""
CareerAgent - AI-Powered Job Assistant refactored to use LangChain framework
Created for Quira Quest 25 by Chaitanya Sharma
"""

import os
import fitz  # PyMuPDF
import requests
import streamlit as st
from io import BytesIO, StringIO
import logging
from typing import List, Dict, Any
import tempfile
from datetime import datetime
from docx import Document
import pdfkit
import re
import time

# LangChain imports
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain, SequentialChain
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.tools import Tool
from langchain.agents import AgentExecutor, initialize_agent, AgentType
from langchain.memory import ChatMessageHistory, ConversationBufferMemory
from langchain.output_parsers import CommaSeparatedListOutputParser
from langchain.tools.base import ToolException
from langchain.schema import SystemMessage, HumanMessage, AIMessage

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Environment variables
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
JOOBLE_API_KEY = os.getenv("JOOBLE_API_KEY")

# Initialize LLM
try:
    llm = ChatGoogleGenerativeAI(
        model="gemini-2.0-flash-exp",
        google_api_key=GOOGLE_API_KEY,
        temperature=0.5,
    )
    logging.info("Gemini LLM initialized successfully.")
except Exception as e:
    logging.error(f"Error initializing Gemini LLM: {e}")
    st.error("Failed to initialize LLM. Check your API key.")

def extract_text_from_pdf(file):
    """Extracts text from a PDF file using PyMuPDF."""
    try:
        doc = fitz.open(stream=file.read(), filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        logging.info("Text extracted from PDF successfully.")
        return text
    except Exception as e:
        logging.error(f"Error extracting text from PDF: {e}")
        raise ToolException(f"Error processing PDF: {e}")

def fetch_job_listings(query: str) -> List[Dict[str, str]]:
    """Fetches job listings from the Jooble API."""
    if not query:
        raise ToolException("Empty job search query provided.")

    # Ensure the query is URL-safe
    query = query.strip().replace(" ", "%20")
    url = f"https://jooble.org/api/{JOOBLE_API_KEY}"

    payload = {
        "keywords": query,
        "location": "",
        "radius": "25",
        "page": 1
    }

    try:
        response = requests.post(url, json=payload)
        response.raise_for_status()
        data = response.json()
        jobs = data.get("jobs", [])

        # Validate and clean up job data
        cleaned_jobs = []
        for job in jobs:
            cleaned_jobs.append({
                "title": job.get("title", "Unknown Title"),
                "company_name": job.get("company", "Unknown Company"),
                "location": job.get("location", "Not Specified"),
                "description": job.get("snippet", "No description available."),
            })

        logging.info(f"Fetched {len(cleaned_jobs)} job listings from Jooble API.")
        return cleaned_jobs

    except Exception as e:
        logging.error(f"Error fetching job listings: {e}")
        raise ToolException(f"Failed to fetch job listings: {e}")

class CareerTools:
    @staticmethod
    def setup_tools(resume_text: str = None):
        """Set up LangChain tools for the CareerAgent."""
        
        tools = [
            Tool(
                name="analyze_resume",
                description="Analyze a resume and extract key information like skills, experience, education, etc.",
                func=lambda query: CareerTools.analyze_resume(resume_text, query),
            ),
            Tool(
                name="suggest_job_titles",
                description="Suggest job titles based on resume content",
                func=lambda _: CareerTools.suggest_job_titles(resume_text),
            ),
            Tool(
                name="extract_keywords_from_resume",
                description="Extract relevant keywords from resume for job searching",
                func=lambda _: CareerTools.extract_keywords(resume_text),
            ),
            Tool(
                name="search_jobs",
                description="Search for job listings by keyword or job title",
                func=lambda query: CareerTools.search_jobs(query),
            ),
            Tool(
                name="generate_cover_letter",
                description="Generate a cover letter for a specific job using the resume",
                func=lambda job_info: CareerTools.generate_cover_letter(resume_text, job_info),
            ),
        ]
        return tools
    
    @staticmethod
    def analyze_resume(resume_text: str, query: str) -> str:
        """Analyze the resume based on a specific query."""
        if not resume_text:
            return "No resume uploaded. Please upload a resume first."
            
        prompt = PromptTemplate(
            template="""
            Analyze the following resume based on the query: {query}
            
            RESUME:
            {resume_text}
            
            Provide a detailed analysis addressing the query.
            """,
            input_variables=["resume_text", "query"]
        )
        
        chain = LLMChain(llm=llm, prompt=prompt)
        return chain.run(resume_text=resume_text, query=query)
    
    @staticmethod
    def suggest_job_titles(resume_text: str) -> str:
        """Suggest job titles based on resume content."""
        if not resume_text:
            return "No resume uploaded. Please upload a resume first."
            
        prompt = PromptTemplate(
            template="""
            Based on the following resume, suggest the top 5 job titles that would be a good fit, 
            ranked from most suitable to slightly broader:

            RESUME:
            {resume_text}

            Format the response as a numbered list with brief explanations for why each job title is suitable.
            """,
            input_variables=["resume_text"]
        )
        
        chain = LLMChain(llm=llm, prompt=prompt)
        return chain.run(resume_text=resume_text)
    
    @staticmethod
    def extract_keywords(resume_text: str) -> List[str]:
        """Extract relevant keywords from resume for job searching."""
        if not resume_text:
            return []
            
        output_parser = CommaSeparatedListOutputParser()
        format_instructions = output_parser.get_format_instructions()
        
        prompt = PromptTemplate(
            template="""
            Extract the most relevant keywords for job searching from this resume. 
            Focus on skills, job titles, and technical expertise.
            Return ONLY the top 5 most relevant keywords.
            
            RESUME:
            {resume_text}
            
            {format_instructions}
            """,
            input_variables=["resume_text"],
            partial_variables={"format_instructions": format_instructions}
        )
        
        chain = LLMChain(llm=llm, prompt=prompt)
        result = chain.run(resume_text=resume_text)
        
        try:
            return output_parser.parse(result)
        except Exception as e:
            logging.error(f"Error parsing keywords: {e}")
            # Fallback parsing if output parser fails
            return [keyword.strip() for keyword in result.split(',')[:5]]
    
    @staticmethod
    def search_jobs(query: str) -> Dict:
        """Search for job listings by keyword or job title."""
        try:
            jobs = fetch_job_listings(query)
            
            if not jobs:
                return {"status": "error", "message": "No jobs found matching your query. Try different keywords or broader terms.", "jobs": []}
            
            # Return the full job details in a dict for easy processing
            result = {
                "status": "success",
                "message": f"Found {len(jobs)} job(s) matching your query.",
                "jobs": jobs
            }
            
            return result
            
        except ToolException as e:
            return {"status": "error", "message": f"Error searching for jobs: {str(e)}", "jobs": []}
    
    @staticmethod
    def generate_cover_letter(resume_text: str, job_info: str) -> str:
        """Generate a cover letter for a specific job using the resume."""
        if not resume_text:
            return "No resume uploaded. Please upload a resume first."
            
        try:
            # Parse job info from input - expecting format: "title: X, company: Y, description: Z"
            job_parts = job_info.split(", ")
            job_dict = {}
            
            for part in job_parts:
                if ": " in part:
                    key, value = part.split(": ", 1)
                    job_dict[key.strip()] = value.strip()
            
            job_title = job_dict.get("title", "")
            job_company = job_dict.get("company", "")
            job_description = job_dict.get("description", "")
            
            if not (job_title and job_description):
                return "Please provide both job title and description to generate a cover letter."
            
            prompt = PromptTemplate(
                template="""
                You are a professional career advisor. Write a tailored cover letter for a job application.
                
                RESUME:
                {resume_text}
                
                JOB DETAILS:
                Title: {job_title}
                Company: {job_company}
                Description: {job_description}
                
                Write a professional cover letter that highlights the candidate's relevant experience and skills
                for this specific job. Format it as a proper business letter.
                """,
                input_variables=["resume_text", "job_title", "job_company", "job_description"]
            )
            
            chain = LLMChain(llm=llm, prompt=prompt)
            return chain.run(
                resume_text=resume_text, 
                job_title=job_title, 
                job_company=job_company, 
                job_description=job_description
            )
            
        except Exception as e:
            logging.error(f"Error generating cover letter: {e}")
            return f"Error generating cover letter: {str(e)}"

def save_as_pdf(content: str, job_title: str, company_name: str) -> str:
    """Saves content as a PDF file."""
    try:
        # Clean filename to avoid special characters
        safe_title = re.sub(r'[\\/*?:"<>|]', "", job_title)
        safe_company = re.sub(r'[\\/*?:"<>|]', "", company_name)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Create a temporary HTML file
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>Cover Letter - {safe_title} at {safe_company}</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 1in; line-height: 1.5; }}
                .header {{ margin-bottom: 20px; }}
                .content {{ margin-bottom: 20px; white-space: pre-wrap; }}
                .signature {{ margin-top: 40px; }}
            </style>
        </head>
        <body>
            <div class="content">{content}</div>
        </body>
        </html>
        """
        
        # Create directory if it doesn't exist
        output_dir = os.path.join(os.getcwd(), "cover_letters")
        os.makedirs(output_dir, exist_ok=True)
        
        # Generate file paths
        html_path = os.path.join(output_dir, f"cover_letter_{safe_title}_{safe_company}_{timestamp}.html")
        pdf_path = os.path.join(output_dir, f"cover_letter_{safe_title}_{safe_company}_{timestamp}.pdf")
        
        # Write HTML file
        with open(html_path, "w", encoding="utf-8") as f:
            f.write(html_content)
        
        # Convert to PDF
        try:
            pdfkit.from_file(html_path, pdf_path)
            os.remove(html_path)  # Clean up HTML file
            logging.info(f"Cover letter saved as PDF: {pdf_path}")
            return pdf_path
        except Exception as e:
            logging.error(f"Error converting to PDF, falling back to HTML: {e}")
            return html_path
            
    except Exception as e:
        logging.error(f"Error saving cover letter as PDF: {e}")
        raise ToolException(f"Failed to save cover letter as PDF: {e}")

def save_as_docx(content: str, job_title: str, company_name: str) -> str:
    """Saves content as a DOCX file."""
    try:
        # Clean filename to avoid special characters
        safe_title = re.sub(r'[\\/*?:"<>|]', "", job_title)
        safe_company = re.sub(r'[\\/*?:"<>|]', "", company_name)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Create directory if it doesn't exist
        output_dir = os.path.join(os.getcwd(), "cover_letters")
        os.makedirs(output_dir, exist_ok=True)
        
        # Generate file path
        docx_path = os.path.join(output_dir, f"cover_letter_{safe_title}_{safe_company}_{timestamp}.docx")
        
        # Create document
        doc = Document()
        
        # Add title
        doc.add_heading(f"Cover Letter - {job_title} at {company_name}", 1)
        
        # Add content (split by paragraphs)
        paragraphs = content.split('\n')
        for para in paragraphs:
            if para.strip():  # Skip empty lines
                doc.add_paragraph(para)
        
        # Save document
        doc.save(docx_path)
        logging.info(f"Cover letter saved as DOCX: {docx_path}")
        return docx_path
        
    except Exception as e:
        logging.error(f"Error saving cover letter as DOCX: {e}")
        raise ToolException(f"Failed to save cover letter as DOCX: {e}")

def create_career_agent(resume_text: str = None):
    """Create the LangChain agent for career assistance."""
    
    tools = CareerTools.setup_tools(resume_text)
    
    agent = initialize_agent(
        tools=tools,
        llm=llm,
        agent=AgentType.CHAT_CONVERSATIONAL_REACT_DESCRIPTION,
        verbose=True,
        handle_parsing_errors=True,
        max_iterations=3,
        memory=ConversationBufferMemory(memory_key="chat_history", return_messages=True)
    )
    
    return agent

class ConversationManager:
    def __init__(self):
        self.history = ChatMessageHistory()
        self.system_message = SystemMessage(content="""
        You are CareerAgent, an AI-powered job assistant designed to help people with their job search.
        
        You can:
        1. Analyze resumes to identify key skills and experience
        2. Suggest relevant job titles based on resume content
        3. Extract keywords from resumes for job searching
        4. Search for job listings using keywords or job titles
        5. Generate tailored cover letters for specific jobs
        
        Remember to be helpful, informative, and guide the user through their job search process.
        """)
    
    def add_user_message(self, message: str):
        self.history.add_user_message(message)
    
    def add_ai_message(self, message: str):
        self.history.add_ai_message(message)
    
    def get_chat_history(self):
        return self.history.messages
    
    def format_history_as_string(self):
        messages = self.get_chat_history()
        formatted_history = ""
        
        for message in messages:
            if isinstance(message, HumanMessage):
                formatted_history += f"Human: {message.content}\n"
            elif isinstance(message, AIMessage):
                formatted_history += f"AI: {message.content}\n"
        
        return formatted_history

def main():
    st.set_page_config(page_title="CareerAgent", page_icon="💼", layout="centered")
    st.title("💼 CareerAgent — AI-Powered Job Assistant")
    st.markdown("*Created for Quira Quest 25 by Chaitanya Sharma using LangChain*")
    
    # Session state initialization
    if "resume_text" not in st.session_state:
        st.session_state.resume_text = None
    if "agent" not in st.session_state:
        st.session_state.agent = None
    if "conversation_manager" not in st.session_state:
        st.session_state.conversation_manager = ConversationManager()
    
    # Upload Resume
    resume_file = st.file_uploader("📄 Upload your resume (PDF)", type=["pdf"])
    
    if resume_file:
        with st.spinner("🔍 Processing resume..."):
            try:
                st.session_state.resume_text = extract_text_from_pdf(resume_file)
                st.session_state.agent = create_career_agent(st.session_state.resume_text)
                st.success("✅ Resume processed successfully!")
                
            except Exception as e:
                st.error(f"❌ Error processing resume: {e}")
    
    # Display additional UI elements if resume is processed
    if st.session_state.resume_text:
        # Add helpful information after resume upload instead of chat
        st.info("✅ Your resume is ready for automated job search! Use the section below to find matching jobs and generate cover letters.")
    
        # Add a new automated job search section - keep this section as is since it's the automated part
        st.header("🤖 Automated Job Search")
        st.markdown("Search for jobs, generate cover letters, and save them as PDF or DOCX files.")
        
        # Job search form
        with st.form("job_search_form"):
            col1, col2 = st.columns([2, 1])
            
            with col1:
                job_query = st.text_input("Job Title or Keywords", 
                                         placeholder="Enter job title or keywords", 
                                         key="auto_job_search")
                                         
            with col2:
                use_resume_keywords = st.checkbox("Use Resume Keywords", value=True, 
                                                help="Auto-extract and include keywords from your resume")
                
            submit_search = st.form_submit_button("🔍 Search Jobs")
        
        # Process job search
        if submit_search and (job_query or use_resume_keywords):
            with st.spinner("🔍 Searching for jobs..."):
                # Combine manual query with resume keywords if selected
                search_query = job_query
                
                if use_resume_keywords:
                    resume_keywords = CareerTools.extract_keywords(st.session_state.resume_text)
                    # Add top keywords if they have a manual query, otherwise use all keywords
                    if job_query:
                        keyword_str = " ".join(resume_keywords[:2])  # Use just top 2 keywords with manual query
                        search_query = f"{job_query} {keyword_str}"
                    else:
                        search_query = " ".join(resume_keywords)
                        
                # Store the search query in session state to display it
                st.session_state.current_search_query = search_query
                
                # Get search results
                search_result = CareerTools.search_jobs(search_query)
                
                # Store jobs in session state
                if search_result["status"] == "success":
                    st.session_state.job_search_results = search_result["jobs"]
                    st.session_state.job_search_message = search_result["message"]
                else:
                    st.session_state.job_search_results = []
                    st.session_state.job_search_message = search_result["message"]
        
        # Display search results if available
        if hasattr(st.session_state, 'job_search_results'):
            if st.session_state.job_search_results:
                st.success(st.session_state.job_search_message)
                st.markdown(f"🔍 Search query: **{st.session_state.current_search_query}**")
                
                # Create tabs for each job
                job_tabs = st.tabs([f"{job['title']} at {job['company_name']}" 
                                   for job in st.session_state.job_search_results[:10]])
                
                # Display each job in a tab with cover letter generation and saving options
                for i, tab in enumerate(job_tabs):
                    job = st.session_state.job_search_results[i]
                    
                    with tab:
                        st.markdown(f"### {job['title']}")
                        st.markdown(f"**Company:** {job['company_name']}")
                        st.markdown(f"**Location:** {job['location']}")
                        st.markdown("**Description:**")
                        st.markdown(job['description'])
                        
                        # Cover letter section
                        st.divider()
                        st.subheader("📝 Cover Letter")
                        
                        col1, col2 = st.columns([1, 2])
                        
                        with col1:
                            if st.button("Generate Cover Letter", key=f"gen_{i}"):
                                with st.spinner("✍️ Writing your tailored cover letter..."):
                                    job_info = f"title: {job['title']}, company: {job['company_name']}, description: {job['description']}"
                                    cover_letter = CareerTools.generate_cover_letter(st.session_state.resume_text, job_info)
                                    # Store in session state with job index as key
                                    st.session_state[f"cover_letter_{i}"] = cover_letter
                        
                        # If cover letter exists for this job, show it and add download options
                        if f"cover_letter_{i}" in st.session_state:
                            st.text_area("Cover Letter", st.session_state[f"cover_letter_{i}"], height=300)
                            
                            # Download options
                            st.subheader("💾 Save Cover Letter")
                            col1, col2 = st.columns(2)
                            
                            with col1:
                                if st.button("Save as PDF", key=f"pdf_{i}"):
                                    with st.spinner("Creating PDF..."):
                                        try:
                                            pdf_path = save_as_pdf(
                                                st.session_state[f"cover_letter_{i}"],
                                                job['title'],
                                                job['company_name']
                                            )
                                            st.success(f"✅ Cover letter saved as PDF: {pdf_path}")
                                        except Exception as e:
                                            st.error(f"Error saving PDF: {str(e)}")
                            
                            with col2:
                                if st.button("Save as DOCX", key=f"docx_{i}"):
                                    with st.spinner("Creating DOCX..."):
                                        try:
                                            docx_path = save_as_docx(
                                                st.session_state[f"cover_letter_{i}"],
                                                job['title'],
                                                job['company_name']
                                            )
                                            st.success(f"✅ Cover letter saved as DOCX: {docx_path}")
                                        except Exception as e:
                                            st.error(f"Error saving DOCX: {str(e)}")
            else:
                st.warning(st.session_state.job_search_message)
        
        st.divider()
        
        # Original Quick Actions section
        with st.expander("Quick Actions"):
            col1, col2, col3 = st.columns(3)
            
            with col1:
                if st.button("📊 Suggest Job Titles"):
                    with st.spinner("Generating suggestions..."):
                        result = CareerTools.suggest_job_titles(st.session_state.resume_text)
                        st.markdown("### Suggested Job Titles")
                        st.markdown(result)
                        st.session_state.job_titles_result = result
            
            with col2:
                if st.button("🔍 Extract Keywords"):
                    with st.spinner("Extracting keywords..."):
                        keywords = CareerTools.extract_keywords(st.session_state.resume_text)
                        result = "Based on your resume, here are the most relevant keywords for your job search:\n\n" + \
                                 ", ".join(keywords)
                        st.markdown("### Resume Keywords")
                        st.markdown(result)
                        st.session_state.keywords_result = result
            
            with col3:
                job_search = st.text_input("🔎 Quick Job Search", placeholder="Enter job title")
                if job_search and job_search != st.session_state.get('last_job_search', ''):
                    with st.spinner("Searching jobs..."):
                        search_result = CareerTools.search_jobs(job_search)
                        st.session_state.last_job_search = job_search
                        
                        # Store search results in session state
                        if search_result["status"] == "success":
                            st.session_state.job_search_results = search_result["jobs"]
                            st.session_state.job_search_message = search_result["message"]
                            st.session_state.current_search_query = job_search
                        else:
                            st.session_state.job_search_results = []
                            st.session_state.job_search_message = search_result["message"]
                        
                        # Trigger rerun to display the results
                        st.session_state.rerun_triggered = True

    # Add a controlled rerun at the end of the main function if triggered
    if st.session_state.get('rerun_triggered', False):
        st.session_state.rerun_triggered = False
        st.rerun()

if __name__ == "__main__":
    main()